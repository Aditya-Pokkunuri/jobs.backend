"""
Utility to convert Markdown text into a .docx (Word) document.
Used by the Resume Builder download endpoint.
"""

import re
from io import BytesIO

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def generate_docx_from_markdown(markdown_text: str) -> BytesIO:
    """
    Convert a Markdown string into a .docx file returned as a BytesIO stream.

    Supports:
      - Headings: #, ##, ###
      - Bold: **text**
      - Bullet points: lines starting with - or *
      - Regular paragraphs
    """
    doc = Document()

    # ── Default style tweaks ─────────────────────────────────
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    lines = markdown_text.split("\n")

    for line in lines:
        stripped = line.strip()

        # Skip empty lines
        if not stripped:
            continue

        # ── Headings ─────────────────────────────────────────
        if stripped.startswith("### "):
            p = doc.add_heading(stripped[4:], level=3)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            continue

        if stripped.startswith("## "):
            p = doc.add_heading(stripped[3:], level=2)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            continue

        if stripped.startswith("# "):
            p = doc.add_heading(stripped[2:], level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            continue

        # ── Bullet points ────────────────────────────────────
        if stripped.startswith("- ") or stripped.startswith("* "):
            bullet_text = stripped[2:]
            p = doc.add_paragraph(style="List Bullet")
            _add_formatted_text(p, bullet_text)
            continue

        # ── Horizontal rule (skip) ───────────────────────────
        if stripped in ("---", "***", "___"):
            continue

        # ── Regular paragraph ────────────────────────────────
        p = doc.add_paragraph()
        _add_formatted_text(p, stripped)

    # Write to BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def _add_formatted_text(paragraph, text: str) -> None:
    """
    Parse inline Markdown bold (**text**) and add runs to the paragraph.
    """
    # Split on bold markers
    parts = re.split(r"(\*\*.*?\*\*)", text)

    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)
